import sys
sys.stdout.reconfigure(encoding='utf-8')
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()

section = doc.sections[0]
section.top_margin = Cm(2.0)
section.bottom_margin = Cm(2.0)
section.left_margin = Cm(2.5)
section.right_margin = Cm(2.5)

def heading(doc, text, level=1):
    return doc.add_heading(text, level=level)

def add_para(doc, text, bold=False, italic=False, size=11, color=None, space_before=0, space_after=6):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = RGBColor(*color)
    return p

def add_bullet(doc, bold_part=None, rest=None, text=None):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_after = Pt(3)
    if bold_part:
        r1 = p.add_run(bold_part)
        r1.font.bold = True
        r1.font.size = Pt(11)
        r2 = p.add_run(rest or '')
        r2.font.size = Pt(11)
    else:
        r = p.add_run(text or '')
        r.font.size = Pt(11)
    return p

def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def add_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    hdr_row = table.rows[0]
    for i, h in enumerate(headers):
        cell = hdr_row.cells[i]
        cell.paragraphs[0].clear()
        run = cell.paragraphs[0].add_run(h)
        run.font.bold = True
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(255, 255, 255)
        cell.paragraphs[0].paragraph_format.space_after = Pt(0)
        set_cell_bg(cell, '1F4E79')
    for ri, row_data in enumerate(rows):
        row = table.rows[ri + 1]
        for ci, val in enumerate(row_data):
            cell = row.cells[ci]
            cell.paragraphs[0].clear()
            run = cell.paragraphs[0].add_run(str(val))
            run.font.size = Pt(10)
            cell.paragraphs[0].paragraph_format.space_after = Pt(0)
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(w)
    return table

def add_speaker_box(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.left_indent = Cm(0.5)
    p.paragraph_format.right_indent = Cm(0.5)
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), 'EBF3FB')
    pPr.append(shd)
    r1 = p.add_run('CO RICT: ')
    r1.font.bold = True
    r1.font.size = Pt(10)
    r1.font.color.rgb = RGBColor(31, 78, 121)
    r2 = p.add_run(text)
    r2.font.size = Pt(10)
    r2.font.italic = True
    return p

def separator(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run('_' * 100)
    r.font.size = Pt(7)
    r.font.color.rgb = RGBColor(180, 180, 180)

# ============================================================
# TITLE
# ============================================================
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(10)
r = p.add_run('ML2 - CAST 4: MODELOVANI A HYPERPARAMETER TUNING')
r.font.size = Pt(18)
r.font.bold = True
r.font.color.rgb = RGBColor(31, 78, 121)

p2 = doc.add_paragraph()
p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
r2 = p2.add_run('Pruvodce prezentaci + odborne vysvetleni modelu a hyperparametru')
r2.font.size = Pt(12)
r2.font.italic = True

p3 = doc.add_paragraph()
p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
r3 = p3.add_run('Analyza studentskeho dusevniho zdravi | VSE | 2026')
r3.font.size = Pt(11)
r3.font.color.rgb = RGBColor(100, 100, 100)

separator(doc)

# ============================================================
# 0. UVOD
# ============================================================
heading(doc, '0. Co budes prezentovat - rychly prehled', level=1)

add_speaker_box(doc,
    '"Dobry den, jmenuji se [jmeno] a predstavim ctvrtu cast nasi prace - modelovani a '
    'optimalizaci hyperparametru. Cil byl natrenovatt pet klasifikacnich modelu, ktere '
    'predikuji, zda student trpi depresemi, a najit pro kazdy model nejlepsi nastaveni. '
    'Zacnu principem kazdeho modelu, pak popisu, jake hyperparametry jsme ladili a proc, '
    'a zakoncim srovanim vysledku."'
)

add_para(doc, 'Vase cast zahrnuje:', bold=True)
add_bullet(doc, text='Jak funguje kazdy z 5 modelu - princip prace')
add_bullet(doc, text='Proc jsme ladili prave tyto hyperparametry a jake hodnoty zvitezily')
add_bullet(doc, text='Jak jsme hledali nejlepsi kombinaci (GridSearch, RandomizedSearch, Optuna)')
add_bullet(doc, text='Srovnani vysledku vsech modelu na testovacich datech')

separator(doc)

# ============================================================
# 1. PREHLED MODELU
# ============================================================
heading(doc, '1. Prehled vsech 5 modelu', level=1)

add_para(doc,
    'Trenovali jsme celkem 5 modelu. Kazdy reprezentuje jiny pristup ke klasifikaci - '
    'od trivialniho baseline az po pokrocily gradient boosting s Bayesovskou optimalizaci.',
    size=11
)

add_table(doc,
    ['Model', 'Metoda hledani parametru', 'Pocet kombinaci', 'ROC-AUC (CV)'],
    [
        ['Baseline (DummyClassifier)', '--- zadne ladeni ---', '---', '0.508'],
        ['Logisticka regrese', 'GridSearchCV', '12', '0.9218'],
        ['Rozhodovaci strom', 'GridSearchCV', '1 440', '0.9040'],
        ['Nahodny les (Random Forest)', 'RandomizedSearchCV (60 iter.)', '60', '0.916'],
        ['XGBoost', 'Optuna - Bayesovska opt. (50 iter.)', '50', '0.9202'],
    ],
    col_widths=[5.0, 5.5, 3.0, 3.0]
)

doc.add_paragraph()
add_speaker_box(doc,
    '"Ke kazdemu modelu jsme hledali optimalni hyperparametry. Ruzne modely jsme ladili '
    'ruznou metodou podle toho, jak velky prostor parametru bylo potreba prohledat. '
    'Vse probihalo vyhradne na trenovacich datech - testovaci mnozina zustala az do evaluace nevizena."'
)

separator(doc)

# ============================================================
# 2. BASELINE
# ============================================================
heading(doc, '2. Model 1 - Baseline (DummyClassifier)', level=1)

heading(doc, 'Princip fungovani', level=2)
add_para(doc,
    'DummyClassifier neprovadi zadne skutecne uceni. Predikuje tridy nahodne '
    'podle jejich cetnosti v trenovacich datech. V nasi sade bylo priblizne 59 % '
    'studentu s depresemi, takze model priradi tridu "deprese" s pravdepodobnosti 59 %. '
    'Vysledek je ciste nahodny a nezavisly na vstupnich priznacich.',
    size=11
)
add_bullet(doc, bold_part='Ucel:', rest=' dolni hranice vykonu - kazdy smysluplny model musi Baseline prekonat')
add_bullet(doc, bold_part='Vysledek na test. datech:', rest=' Accuracy 52.2 %, ROC-AUC 0.508, celkove naklady 7 978')
add_bullet(doc, bold_part='Hyperparametry:', rest=' zadne k ladeni')

add_speaker_box(doc,
    '"Baseline nam rika: kdybychom nehledali zadne vzory a jen nahodne tipovali '
    'podle pomeru trid, dostali bychom ROC-AUC 0.51 - temer nahodu. Vsechny nase '
    'realne modely musi tuto hodnotu vyrazne prekonat."'
)

separator(doc)

# ============================================================
# 3. LOGISTICKA REGRESE
# ============================================================
heading(doc, '3. Model 2 - Logisticka regrese', level=1)

heading(doc, 'Princip fungovani', level=2)
add_para(doc,
    'Logisticka regrese je linearni klasifikacni model. Predikuje pravdepodobnost '
    'prislusnosti ke tride pomoci sigmoid (logisticke) funkce:',
    size=11
)
add_para(doc, '    P(deprese) = 1 / (1 + e^(-(w0 + w1*x1 + w2*x2 + ...)))', size=11, italic=True)
add_para(doc,
    'Model se uci vahy (w) pro kazdy priznak - napriklad, jak silne "Academic Pressure" '
    'nebo "Suicidal Thoughts" prispivaji k pravdepodobnosti deprese. Vystup je vzdy cislo mezi 0 a 1.',
    size=11
)
add_bullet(doc, bold_part='Vyhody:', rest=' rychly, interpretovatelny - kazdy koeficient ma jasny vyznam')
add_bullet(doc, bold_part='Nevyhody:', rest=' predpoklada linearni vztahy - nezachyti slozite interakce')
add_bullet(doc, bold_part='Regularizace:', rest=' bez ni by model priradeloval velke vahy mene dulezitym priznakum (preuiceni)')

heading(doc, 'Hyperparametry - co ladime a proc', level=2)

add_table(doc,
    ['Hyperparametr', 'Co rika', 'Hodnoty v gridu', 'Vitez', 'Zduvodneni'],
    [
        ['C',
         'Sila regularizace. Male C = silna penalizace velkych vah. Velke C = slaba penalizace.',
         '0.001 / 0.01 / 0.1 / 1 / 10 / 100',
         'C = 0.1',
         'Mirna regularizace. Prilis velke C => preuiceni (model si pamatuje trenovaci data). Prilis male C => prilis jednoduchy model. C=0.1 je optimum.'],
        ['l1_ratio (penalty)',
         'L1 (Lasso) muze vynulovat cely priznak. L2 (Ridge) snizi vsechny vahy rovnomerne.',
         '0.0 = L2, 1.0 = L1',
         'l1_ratio = 1.0 (L1)',
         'L1 provadi automaticky vyber priznaku - mene dulezite priznaky dostaji nulovou vahu. Vhodne po feature selection.'],
    ],
    col_widths=[3.0, 4.5, 3.2, 2.0, 5.3]
)

doc.add_paragraph()
add_para(doc, 'Metoda ladeni: GridSearchCV, 5-nasobna stratifikovana krizova validace', bold=True, size=11)
add_bullet(doc, text='Prohledano: 12 kombinaci (6 hodnot C x 2 typy penalty)')
add_bullet(doc, text='Nejlepsi ROC-AUC na CV: 0.9218')
add_bullet(doc, text='Vysledek na test. datech: Accuracy 84.3 %, ROC-AUC 0.919, Recall 84.8 %, celkove naklady 2 870')

add_speaker_box(doc,
    '"U logisticke regrese jsme ladili hlavne silu regularizace C. Prilis velke C = slaba '
    'penalizace, model se muze preuct. Prilis male C = prilis zjednoduseny model. '
    'Nejlepsi C=0.1 s L1 penalizaci znamena, ze model sam vyhodil mene dulezite priznaky '
    '- L1 regularizace totiz dokaze vahy uplne vynulovat, cimz funguje jako vestavenej '
    'vyber priznaku. Vysledek je ROC-AUC 0.919, coz je druhy nejlepsi ze vsech modelu."'
)

separator(doc)

# ============================================================
# 4. ROZHODOVACI STROM
# ============================================================
heading(doc, '4. Model 3 - Rozhodovaci strom (Decision Tree)', level=1)

heading(doc, 'Princip fungovani', level=2)
add_para(doc,
    'Rozhodovaci strom rozdeli data pomoci serie podminek - podobne jako vyvojovy diagram. '
    'V kazdem uzlu algoritmus hleda priznak a prahovou hodnotu, ktera nejlepe rozdeli '
    'vzorky na dve skupiny. Cistotu uzlu meri Gini index nebo entropie.',
    size=11
)
add_para(doc, 'Priklad z naseho stromu (prvni 3 urovne):', bold=True, size=11)
add_bullet(doc, text='Koren: Mel student sebevrazedne myslenky?')
add_bullet(doc, text='  -> ANO  =>  pravdepodobnost deprese je velmi vysoka')
add_bullet(doc, text='  -> NE   =>  Stress_Index <= 3.25?')
add_bullet(doc, text='        -> Nizky stres  =>  Age <= 29  =>  ...')
add_para(doc,
    'Gini impurity: mira necistoty uzlu. Hodnota 0 = pouze jedna trida v uzlu (idealni). '
    'Hodnota 0.5 = dokonala nerozhodnost (50/50 mix trid).',
    italic=True, size=10
)
add_bullet(doc, bold_part='Vyhody:', rest=' vizualne interpretovatelny, zadne skaloovani dat potreba')
add_bullet(doc, bold_part='Nevyhody:', rest=' extremne nachylny k preuiceni - bez omezeni roste do hloubky a memoruje trenovaci data')

heading(doc, 'Hyperparametry - co ladime a proc', level=2)

add_table(doc,
    ['Hyperparametr', 'Co rika', 'Hodnoty v gridu', 'Vitez', 'Zduvodneni'],
    [
        ['max_depth', 'Max. hloubka stromu', '3, 5, 8, 12, None', '8',
         'Klic. parametr. Neomezeny strom si zapamatuje kazdy trenovaci priklad. Hloubka 8 nabizi dostatecnou komplexitu bez preuiceni.'],
        ['min_samples_leaf', 'Min. vzorku v kazdem listu stromu', '1, 5, 10, 20', '20',
         'Listy s malo vzorky jsou preuicene. 20 vzorku zajistuje statisticky spolehlivejsi listy a lepsi generalizaci na novych datech.'],
        ['min_samples_split', 'Min. vzorku v uzlu aby se smel rozdelit', '2, 10, 20', '2',
         'V kombinaci s min_samples_leaf (20) je hodnota 2 dostacujici - leaf constraint je silnejsi.'],
        ['max_features', 'Kolik priznaku uvazovat v kazdem uzlu', 'sqrt, log2, None', 'None (vsechny)',
         'S malym poctem priznaku po feature selection je vyhodne uvazovat vsechny - sqrt by zbytecne omezoval.'],
        ['criterion', 'Mira necistoty pro vyber vetveni', 'gini, entropy', 'entropy',
         'Entropie lepe rozlisuje u nevyvazenych trid. Rozdil je obvykle maly, ale entropie zvitezila v CV.'],
        ['ccp_alpha', 'Post-pruning - orez vetvi po trenovani', '0.0, 0.0005, 0.001, 0.005', '0.0005',
         'Jemny orez vetvi s malym prinosem PO trenovani. ccp_alpha=0.0005 zlepsil generalizaci bez ztraty informace.'],
    ],
    col_widths=[3.2, 4.0, 3.5, 2.0, 5.3]
)

doc.add_paragraph()
add_para(doc, 'Metoda ladeni: GridSearchCV - vycerpavajici prohledani', bold=True, size=11)
add_bullet(doc, text='Prohledano: 1 440 kombinaci (5 x 4 x 3 x 3 x 2 x 4)')
add_bullet(doc, text='Nejlepsi ROC-AUC na CV: 0.9040')
add_bullet(doc, text='Vysledek na test. datech: Accuracy 82.2 %, ROC-AUC 0.902, celkove naklady 3 419')

add_speaker_box(doc,
    '"Rozhodovaci strom je bez omezeni extremne nachylny na preuiceni - '
    'neomezeny strom si fakticky zapamatuje kazdy trenovaci priklad. '
    'Proto jsme ladili sest parametru, ktere komplexitu omezuji. '
    'Klicovy je max_depth=8 kombinovany s ccp_alpha=0.0005 jako post-pruning '
    'a min_samples_leaf=20. Zajimave je, ze strom dosahuje horsiho ROC-AUC nez '
    'logisticka regrese - linearni vztahy v datech jsou tak silne, ze jednoduchy '
    'linearni model si vede lepe nez stromova struktura."'
)

separator(doc)

# ============================================================
# 5. RANDOM FOREST
# ============================================================
heading(doc, '5. Model 4 - Nahodny les (Random Forest)', level=1)

heading(doc, 'Princip fungovani', level=2)
add_para(doc,
    'Nahodny les je ensemble metoda - kombinuje predikce stovek rozhodovacich stromu. '
    'Kazdy strom se trenuje na jinem vzorku dat (bootstrap sampling) a v kazdem uzlu '
    'uvazuje jen nahodnou podmnozu priznaku. Finalni predikce je hlasovani majority.',
    size=11
)
add_para(doc, 'Proc je RF lepsi nez jeden strom?', bold=True, size=11)
add_bullet(doc, bold_part='Bias-variance tradeoff:',
           rest=' jeden hluboky strom = nizky bias, vysoky variance (preuiceni). '
                'Ensemble 300 stromu = prumerovani variance vyrazne snizuje.')
add_bullet(doc, bold_part='Nahodnost priznaku:',
           rest=' stromy jsou mene korelovany => vetsi diverzita => lepsi generalizace')
add_bullet(doc, bold_part='Bootstrap sampling:',
           rest=' kazdy strom vidi ~63 % dat (s opakovanim); zbyvajicich 37 % slouzi jako out-of-bag validace')

heading(doc, 'Hyperparametry - co ladime a proc', level=2)

add_table(doc,
    ['Hyperparametr', 'Co rika', 'Rozsah/hodnoty', 'Zduvodneni proc ladime'],
    [
        ['n_estimators', 'Pocet stromu v lese', '100-700 (nahodne)',
         'Vice stromu = stabilnejsi predikce; nad ~300 prinos klesa. Hledame optimum vykon vs. cas vypoctu.'],
        ['max_depth', 'Hloubka kazdeho stromu', '5, 8, 10, 15, 20, None',
         'RF toleruje hlubsi stromy nez single DT, protoze ensemble prumerovani kompenzuje preuiceni.'],
        ['min_samples_leaf', 'Min. vzorku v listu', '1-15 (nahodne)',
         'Kontrola komplexity kazdeho stromu; vyssi hodnota = konzervativnejsi vetveni.'],
        ['min_samples_split', 'Min. vzorku pro vetveni', '2-20 (nahodne)',
         'Spolupusobi s min_samples_leaf - omezuje prilis specificke vetveni na trenovacich datech.'],
        ['max_features', 'Pocet priznaku uvazovanych v kazdem uzlu', 'sqrt, log2, 0.3, 0.5',
         'Klicovy pro dekorelaci stromu. sqrt je standard pro klasifikaci - mensi hodnota = diverzita, ale vyssi bias.'],
        ['bootstrap', 'Vzorkovani s opakovanim', 'True / False',
         'True = klasicky RF s bootstrap samplingem. False = kazdy strom vidi vsechna data (mene rozmanitosti stromu).'],
    ],
    col_widths=[3.2, 4.0, 3.5, 7.3]
)

doc.add_paragraph()
add_para(doc, 'Metoda ladeni: RandomizedSearchCV - 60 nahodnych kombinaci', bold=True, size=11)
add_bullet(doc, text='Proc Randomized misto Grid? Kombinaci je prilis mnoho - vycerpavajici Grid by trval hodiny.')
add_bullet(doc, text='60 nahodnych vzorku z distribuce dobre pokryje prostor parametru.')
add_bullet(doc, text='Nejlepsi ROC-AUC na CV: 0.916')
add_bullet(doc, text='Vysledek na test. datech: Accuracy 83.6 %, ROC-AUC 0.916, celkove naklady 2 930')

add_speaker_box(doc,
    '"Nahodny les je vyrazne robustnejsi nez jeden strom. Kombinuje stovky imperfektnich '
    'stromu a jejich prumerovanim snizuje chybu z preuiceni. Prostor hyperparametru je ale '
    'veliky, proto jsme pouzili RandomizedSearchCV - 60 nahodnych kombinaci je dostacujicich '
    'k nalezeni rozumneho optima bez vycerpavajiciho grid search."'
)

separator(doc)

# ============================================================
# 6. XGBOOST
# ============================================================
heading(doc, '6. Model 5 - XGBoost + Optuna (Bayesovska optimalizace)', level=1)

heading(doc, 'Princip fungovani XGBoost - gradient boosting', level=2)
add_para(doc,
    'XGBoost (eXtreme Gradient Boosting) funguje zasadne jinak nez nahodny les. '
    'Stromy nestavuje paralelne, ale POSTUPNE - kazdy novy strom se snazi opravit '
    'chyby vsech predchozich. Tento pristup se nazyva gradient boosting.',
    size=11
)
add_para(doc, 'Jak to funguje krok za krokem:', bold=True, size=11)
add_bullet(doc, bold_part='Iterace 1:', rest=' trenuje se jednoduchy strom na puvodnich datech')
add_bullet(doc, bold_part='Iterace 2:', rest=' druhy strom se trenuje na REZIDUALECH (chybach) prvniho stromu')
add_bullet(doc, bold_part='Iterace 3+:', rest=' kazdy dalsi strom opravuje zbyvajici chyby predchozich')
add_bullet(doc, bold_part='Vysledek:', rest=' vazeny soucet vsech stromu - kazdy prispiva malou korekcí (learning_rate)')
add_para(doc,
    'XGBoost navic pridava L1 a L2 regularizaci primo do objektivni funkce, '
    'coz ho dela odolnejsim vuci preuiceni nez klasicky Gradient Boosting.',
    size=11, italic=True
)
add_bullet(doc, bold_part='Vyhody pred RF:', rest=' obvykle lepsi presnost, efektivnejsi vyuziti kazdeho vzorku')
add_bullet(doc, bold_part='Nevyhody:', rest=' vice hyperparametru, sekvencni trenovani (pomalejsi), citlivejsi na learning_rate')

heading(doc, 'Hyperparametry XGBoost - co ladime a proc', level=2)

add_table(doc,
    ['Hyperparametr', 'Co rika', 'Rozsah pro Optunu', 'Zduvodneni proc ladime'],
    [
        ['n_estimators', 'Pocet boostovacich iteraci (stromu)', '50 - 500',
         'Vice = lepsi aproximace, ale riziko preuiceni; musi byt v rovnovaze s learning_rate.'],
        ['max_depth', 'Hloubka kazdeho stromu', '2 - 8',
         'XGBoost preferuje melci stromy (3-6) nez RF. Hloubka je kompenzovana poctem iteraci boosting.'],
        ['learning_rate', 'Krok uceni (eta) - jak moc kazdy strom prispiva k finalni predikci', '0.01 - 0.3',
         'Male eta = opatrne uceni, lepsi generalizace, ale potrebuje vice stromu (n_estimators). Klicovy parametr.'],
        ['subsample', 'Podil dat pouzitych pro trenovani kazdeho stromu', '0.5 - 1.0',
         'Nahodny vyber podmnoiny dat na kazdou iteraci - snizuje preuiceni (analogie k bootstrap u RF).'],
        ['colsample_bytree', 'Podil priznaku pouzitych pro kazdy strom', '0.5 - 1.0',
         'Analogie k max_features u RF - dekoreluje stromy a snizuje preuiceni.'],
        ['min_child_weight', 'Min. suma vah vzorku v listu stromu', '1 - 10',
         'Vyssi = konzervativnejsi strom; dulezite pri nevyvazenych tridach aby model nezanasely vzacne vzorky.'],
        ['gamma', 'Min. zisk pro rozdeleni uzlu', '0 - 0.5',
         'Vyssi = "lenivejsi" strom - vetvi se pouze tam, kde je dostatecny prinos. Cistsi struktura stromu.'],
        ['reg_alpha', 'L1 regularizace (Lasso) na listove vahach', '0 - 1',
         'Muze vynulovat prispevky mene dulezitych priznaku - vesstavena feature selection.'],
        ['reg_lambda', 'L2 regularizace (Ridge) na listove vahach', '0.5 - 5',
         'Brzdi velke korekce jednoho stromu - standardni pro XGBoost. Zabrauje dominanci jedne iterace.'],
        ['scale_pos_weight', 'Vaha pozitivni tridy (deprese) pri vypoctu ztraty', 'sum(neg)/sum(pos) ~ 0.71',
         'Kompenzuje nevyvazenost trid - bez toho by model ignoroval mensinovou tridu (studenty bez deprese).'],
    ],
    col_widths=[3.5, 4.5, 3.2, 6.8]
)

doc.add_paragraph()
heading(doc, 'Proc Optuna misto GridSearch nebo RandomizedSearch?', level=2)
add_para(doc,
    'XGBoost ma 10 hyperparametru. GridSearch s rozumnymi hodnotami by znamenal '
    'tisice kombinaci - trval by hodiny nebo dny. RandomizedSearch vzorkuje nahodne, '
    'ale ignoruje vysledky predchozich pokusu.',
    size=11
)
add_para(doc, 'Optuna pouziva Bayesovskou optimalizaci s TPE samplerem:', bold=True, size=11)
add_bullet(doc, text='Po kazde zkousce aktualizuje pravdepodobnostni model - kde v prostoru jsou dobre vysledky')
add_bullet(doc, text='Pristi zkouska se umisti tam, kde je nejvyssi Expected Improvement')
add_bullet(doc, text='50 iteraci Optuny je srovnatelnych efektivitou s 300+ nahodnymi zkouskami')
add_bullet(doc, text='TPE = Tree-structured Parzen Estimator - uci se z predchozich zkousek, ne nahodne')

add_para(doc, 'Vysledky XGBoost:', bold=True, size=11)
add_bullet(doc, text='Nejlepsi ROC-AUC na CV: 0.9202')
add_bullet(doc, text='Vysledek na test. datech: Accuracy 84.5 %, ROC-AUC 0.920, celkove naklady 2 849 (nejlepsi ze vsech!)')

add_speaker_box(doc,
    '"XGBoost je nas nejsilnejsi model - dosahl nejlepsiho ROC-AUC i nejnizsich celkovych '
    'nakladu. Klicovy je parametr learning_rate - male hodnoty 0.01-0.05 davaji lepsi '
    'generalizaci, ale vyzaduji vice stromu. scale_pos_weight automaticky kompenzuje '
    'nevyvazenost - mame vice depresivnich studentu, takze mensinova trida (bez deprese) '
    'by jinak byla ignorovana. Optuna nam umoznila efektivne prohledat velky prostor '
    '10 hyperparametru bez zbytecne vycerpavajiciho grid search."'
)

separator(doc)

# ============================================================
# 7. SROVNANI
# ============================================================
heading(doc, '7. Vysledky - srovnani vsech modelu po tuningu', level=1)

add_para(doc, 'Vysledky na testovacich datech pri standardnim prahu 0.5:', bold=True, size=11)
add_table(doc,
    ['Model', 'Accuracy', 'Recall (deprese)', 'F1', 'ROC-AUC', 'Celkove naklady'],
    [
        ['Baseline', '52.2 %', '59.4 %', '0.593', '0.508', '7 978'],
        ['Logisticka regrese', '84.3 %', '84.8 %', '0.863', '0.919', '2 870'],
        ['Rozhodovaci strom', '82.2 %', '81.5 %', '0.843', '0.902', '3 419'],
        ['Nahodny les', '83.6 %', '84.6 %', '0.858', '0.916', '2 930'],
        ['XGBoost (Optuna)', '84.5 %', '84.8 %', '0.865', '0.920', '2 849 (NEJLEPSI)'],
    ],
    col_widths=[4.5, 2.5, 3.5, 2.0, 2.5, 3.5]
)

doc.add_paragraph()
add_para(doc, 'Klicova zjisteni z tuningu:', bold=True, size=11)
add_bullet(doc, text='Vsechny natrenovane modely vyrazne prekonaly Baseline (ROC-AUC 0.51 => 0.90-0.92)')
add_bullet(doc, text='XGBoost dosahl nejlepsiho vysledku ve vsech metrikach')
add_bullet(doc, text='Rozhodovaci strom si vede hure nez ostatni - linearni vztahy v datech zvyhodruji regresni modely')
add_bullet(doc, text='Tuning prinesl meritelne zlepseni u kazdeho modelu oproti vychozim parametrum')
add_bullet(doc, text='Veskere ladeni probihalo vyhradne na trenovacich datech - testovaci mnozina zustala nevizena')

add_speaker_box(doc,
    '"Tim moje cast prezentace konci. Shrneme: natrenovani jsme 5 modelu, kazdy s jinou '
    'architekturou a metodou ladeni hyperparametru. XGBoost s Bayesovskou optimalizaci '
    'dosahl nejlepsiho vykonu - ROC-AUC 0.920 a nejnizsi celkove naklady 2 849. '
    'Vysledky predam kolegovi, ktery nyni vysvetli, jak jsme vybrali finalni model '
    'a jak si vede po optimalizaci klasifikacniho prahu."'
)

separator(doc)

# ============================================================
# 8. FAQ
# ============================================================
heading(doc, '8. Mozne otazky od komise', level=1)

faq = [
    ('Proc jste pouzili ROC-AUC jako metriku pro tuning, ne Accuracy?',
     'Data jsou nevyvazena - depresivnich studentu je vice. Accuracy muze byt vysoka, '
     'i kdyz model ignoruje mensinovou tridu. ROC-AUC meri diskriminacni schopnost '
     'nezavisle na prahu rozhodovani a je robustni vuci nevyvazenosti trid.'),
    ('Proc 5-nasobna krizova validace a ne 10-nasobna?',
     '5-nasobna CV je rozumny kompromis vypocetni narocnosti a spolehlivosti odhadu. '
     'S 27 000 vzorky je kazdy fold dostatecne velky. 10-nasobna by trvala 2x dele s minimalnim prinosem.'),
    ('Co je to Gini impurity?',
     'Gini = 1 - suma(pi^2). Vyjadruje pravdepodobnost, ze nahodne vybrany vzorek bude '
     'spatne klasifikovan. Cisty uzel (jen jedna trida) ma Gini=0. '
     'Strom hleda vetveni, ktere nejvice snizi prumerne Gini dcerinnych uzlu.'),
    ('Proc learning_rate u XGBoost musi byt male?',
     'Velky learning_rate znamena, ze kazdy strom silne opravuje predchozi - model '
     'konverguje rychle, ale "preskakuje" optima. Male eta (0.01-0.05) pridava '
     'kazdou korekci opatrneji, model generalizuje lepe, ale potrebuje vice iteraci.'),
    ('Co je TPE sampler v Optuně?',
     'Tree-structured Parzen Estimator - po kazde zkousce aktualizuje dve distribuce: '
     'kde byly dobre vysledky l(x) a kde spatne g(x). Pristi bod vzorkuje '
     'z oblasti kde l(x)/g(x) je nejvyssi - tedy nejvyssi Expected Improvement. '
     'Je to forma Bayesovske optimalizace, ktera se uci z predchozich zkousek.'),
    ('Proc jste nezkusili neuronove site?',
     'Na tabularnich datech gradient boosting jako XGBoost typicky prekonava jednoduche '
     'neuronove site. Navic neuronove site jsou hure interpretovatelne - coz je problem '
     'v kontextu zdravotnich dat, kde chceme rozumet rozhodnutim modelu.'),
    ('Co znamena scale_pos_weight a proc je dulezite?',
     'Tento parametr rika XGBoost, kolikrat vice vahovit pozitivni tridu (deprese). '
     'Pouziva se hodnota sum(neg)/sum(pos). Bez tohoto parametru by model penalizoval '
     'chyby obou trid stejne, ale FN (prehlédnute depresivni) jsou 5x nakladnejsi nez FP.'),
]

for q, a in faq:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    rq = p.add_run('Q: ' + q)
    rq.font.bold = True
    rq.font.size = Pt(11)
    rq.font.color.rgb = RGBColor(31, 78, 121)
    p2 = doc.add_paragraph()
    p2.paragraph_format.left_indent = Cm(0.5)
    p2.paragraph_format.space_after = Pt(6)
    ra = p2.add_run('A: ' + a)
    ra.font.size = Pt(10)

separator(doc)

# ============================================================
# 9. TAHAK
# ============================================================
heading(doc, '9. Tahak - klicova cisla pro prezentaci', level=1)

add_table(doc,
    ['Co rict', 'Cislo / hodnota'],
    [
        ['Celkem modelu', '5 (Baseline, LR, DT, RF, XGBoost)'],
        ['Nejlepsi model', 'XGBoost s Optunou'],
        ['ROC-AUC XGBoost (test)', '0.920'],
        ['ROC-AUC Baseline', '0.508 - temer nahoda'],
        ['Kombinaci v Grid DT', '1 440 kombinaci'],
        ['Iteraci Optuna', '50 Bayesovskych iteraci'],
        ['Nejlepsi C pro logistickou regresi', 'C = 0.1 (L1 regularizace)'],
        ['Nejlepsi max_depth pro DT', '8 (s ccp_alpha=0.0005 post-pruning)'],
        ['Celkove naklady XGBoost', '2 849 (nejnizsi ze vsech)'],
        ['Recall deprese XGBoost', '84.8 % (standard. prah 0.5)'],
        ['Pocet vzorku v datasetu', '27 901 indickych studentu'],
        ['Rozdeleni train/test', '80 % trenovaci / 20 % testovaci (stratifikovane)'],
        ['Metrika pro tuning', 'ROC-AUC (robustni vuci nevyvazenosti trid)'],
        ['Cross-validace', '5-nasobna stratifikovana (StratifiedKFold)'],
    ],
    col_widths=[9.0, 9.0]
)

doc.add_paragraph()
add_para(doc,
    'Hodne stesti na prezentaci zitra! Mluvte pomalu, dychejte. '
    'Kdyz neznate odpoved na otazku, je v poradku rict: '
    '"Tuto otazku bychom museli prozkoumat podrobneji" nebo se otocit na spoluziaka.',
    bold=True, size=11, color=(31, 78, 121)
)

doc.save('prezentace_cast4_modelovani.docx')
print('HOTOVO: prezentace_cast4_modelovani.docx')
